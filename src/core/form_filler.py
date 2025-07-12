import re
from typing import Dict, List, Any
from docx import Document
from datetime import datetime
import json

class FormFiller:
    """Fills detected form fields with extracted information"""
    
    def __init__(self):
        """Initialize the form filler with matching rules"""
        # Define keywords for matching field types to extracted data
        self.field_keywords = {
            'name': ['name', 'full name', 'your name', 'applicant name', 'employee name', 'person name'],
            'date': ['date', 'when', 'day', 'dob', 'birth', 'birthday'],
            'email': ['email', 'e-mail', 'electronic mail', 'email address'],
            'phone': ['phone', 'telephone', 'mobile', 'cell', 'contact number', 'tel'],
            'address': ['address', 'street', 'residence', 'location', 'where you live'],
            'city': ['city', 'town', 'municipality'],
            'state': ['state', 'province'],
            'zip': ['zip', 'postal', 'zip code', 'postal code'],
            'ssn': ['ssn', 'social security', 'social security number'],
            'employee_id': ['employee id', 'employee number', 'staff id', 'worker id'],
            'department': ['department', 'dept', 'division', 'unit'],
            'position': ['position', 'title', 'job title', 'role', 'designation'],
            'salary': ['salary', 'wage', 'pay', 'compensation', 'income']
        }
    
    def match_field_to_data(self, field_label: str, extracted_data: Dict, 
                          transcription_text: str) -> Any:
        """
        Match a form field to extracted data using various strategies
        
        Args:
            field_label: The label of the form field
            extracted_data: Entities extracted from transcription
            transcription_text: Full transcription text for context
            
        Returns:
            Matched value or None
        """
        field_label_lower = field_label.lower()
        
        # Try to match based on field type keywords
        for field_type, keywords in self.field_keywords.items():
            if any(keyword in field_label_lower for keyword in keywords):
                return self._get_value_for_field_type(
                    field_type, 
                    field_label, 
                    extracted_data, 
                    transcription_text
                )
        
        # If no keyword match, try context-based matching
        return self._extract_from_context(field_label, transcription_text)
    
    def _get_value_for_field_type(self, field_type: str, field_label: str,
                                 extracted_data: Dict, text: str) -> Any:
        """Get value based on field type"""
        
        if field_type == 'date' and extracted_data.get('dates'):
            # Return the first date found (can be improved with context)
            return extracted_data['dates'][0]
        
        elif field_type == 'email' and extracted_data.get('emails'):
            return extracted_data['emails'][0]
        
        elif field_type == 'phone' and extracted_data.get('phone_numbers'):
            return extracted_data['phone_numbers'][0]
        
        elif field_type == 'name':
            # Extract name from context
            return self._extract_name_from_context(text)
        
        elif field_type in ['address', 'city', 'state', 'zip']:
            # Extract address components
            return self._extract_address_component(field_type, text)
        
        elif field_type in ['ssn', 'employee_id'] and extracted_data.get('numbers'):
            # Look for specific number patterns
            return self._find_specific_number(field_type, extracted_data['numbers'])
        
        elif field_type in ['department', 'position', 'salary']:
            # Extract from context
            return self._extract_job_info(field_type, text)
        
        return None
    
    def _extract_from_context(self, field_label: str, text: str) -> str:
        """Extract value from context using the field label"""
        # Look for patterns like "field_label is X" or "field_label: X"
        patterns = [
            rf"{re.escape(field_label)}\s*(?:is|:|=)\s*([^\n,.]+)",
            rf"(?:my|the|our)\s+{re.escape(field_label)}\s*(?:is|:|=)\s*([^\n,.]+)",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_name_from_context(self, text: str) -> str:
        """Extract name from transcription context"""
        # Look for name patterns
        patterns = [
            r"(?:my name is|I am|this is)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})",
            r"(?:name|full name)\s*(?:is|:)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2})",
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def _extract_address_component(self, component: str, text: str) -> str:
        """Extract specific address component"""
        # This is a simplified version - can be enhanced with NLP
        if component == 'city':
            # Look for city patterns
            pattern = r"(?:city|town)\s*(?:is|:)\s*([A-Za-z\s]+?)(?:,|\.|$)"
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        elif component == 'state':
            # Look for state abbreviations or full names
            states = ['CA', 'NY', 'TX', 'FL', 'IL', 'PA', 'OH', 'GA', 'NC', 'MI']  # Add more
            for state in states:
                if state in text.upper():
                    return state
        
        elif component == 'zip':
            # Look for 5-digit zip codes
            pattern = r'\b\d{5}\b'
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        
        return None
    
    def _find_specific_number(self, number_type: str, numbers: List[str]) -> str:
        """Find specific type of number from list"""
        if number_type == 'ssn':
            # SSN pattern: XXX-XX-XXXX or 9 digits
            for num in numbers:
                if re.match(r'^\d{3}-?\d{2}-?\d{4}$', num):
                    return num
        
        elif number_type == 'employee_id':
            # Employee IDs are often 4-8 digits
            for num in numbers:
                if 4 <= len(num.replace('-', '')) <= 8:
                    return num
        
        return None
    
    def _extract_job_info(self, info_type: str, text: str) -> str:
        """Extract job-related information"""
        patterns = {
            'department': r"(?:department|dept|division)\s*(?:is|:)\s*([A-Za-z\s]+?)(?:,|\.|$)",
            'position': r"(?:position|title|role)\s*(?:is|:)\s*([A-Za-z\s]+?)(?:,|\.|$)",
            'salary': r"(?:salary|pay|compensation)\s*(?:is|:)\s*\$?([0-9,]+)"
        }
        
        pattern = patterns.get(info_type)
        if pattern:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return None
    
    def fill_form(self, docx_path: str, fields: List[Dict], 
                  extracted_data: Dict, transcription_text: str,
                  output_path: str) -> Dict:
        """
        Fill the form with extracted data
        
        Args:
            docx_path: Path to original form
            fields: List of detected fields
            extracted_data: Extracted entities from transcription
            transcription_text: Full transcription text
            output_path: Path to save filled form
            
        Returns:
            Dictionary with filling results
        """
        doc = Document(docx_path)
        filled_fields = []
        unfilled_fields = []
        
        # Process each field
        for field in fields:
            field_label = field['label']
            field_type = field['field_type']
            
            # Try to find matching data
            value = self.match_field_to_data(
                field_label, 
                extracted_data, 
                transcription_text
            )
            
            if value:
                # Fill the field in the document
                success = self._fill_field_in_document(
                    doc, 
                    field, 
                    value
                )
                
                if success:
                    filled_fields.append({
                        'label': field_label,
                        'type': field_type,
                        'value': value
                    })
                else:
                    unfilled_fields.append(field_label)
            else:
                unfilled_fields.append(field_label)
        
        # Save the filled document
        doc.save(output_path)
        
        # Calculate confidence score
        total_fields = len(fields)
        filled_count = len(filled_fields)
        confidence = filled_count / total_fields if total_fields > 0 else 0
        
        return {
            'filled_fields': filled_fields,
            'unfilled_fields': unfilled_fields,
            'total_fields': total_fields,
            'filled_count': filled_count,
            'confidence_score': confidence,
            'output_path': output_path
        }
    
    def _fill_field_in_document(self, doc: Document, field: Dict, value: str) -> bool:
        """
        Fill a specific field in the document
        
        Args:
            doc: Document object
            field: Field information
            value: Value to fill
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Handle paragraph-based fields
            if isinstance(field.get('paragraph_idx'), int):
                paragraph = doc.paragraphs[field['paragraph_idx']]
                original_text = field['original_text']
                
                # Replace underscores or brackets with the value
                if '___' in original_text:
                    new_text = re.sub(r'_{3,}', value, original_text, count=1)
                elif '[  ]' in original_text or '[ ]' in original_text:
                    # For checkboxes, mark as checked
                    if field['field_type'] == 'checkbox':
                        new_text = original_text.replace('[ ]', '[X]').replace('[  ]', '[X]')
                    else:
                        new_text = re.sub(r'\[[\s]*\]', value, original_text, count=1)
                else:
                    # Try to append value after label
                    new_text = original_text + ' ' + value
                
                # Update paragraph text
                paragraph.text = new_text
                return True
            
            # Handle table-based fields
            elif 'table_' in str(field.get('paragraph_idx', '')):
                # Parse table location
                location = field['paragraph_idx']
                match = re.match(r'table_(\d+)_row_(\d+)', location)
                if match:
                    table_idx = int(match.group(1))
                    row_idx = int(match.group(2))
                    cell_idx = field['run_idx']
                    
                    # Update table cell
                    table = doc.tables[table_idx]
                    cell = table.rows[row_idx].cells[cell_idx]
                    
                    # Replace underscores or append value
                    if '___' in cell.text:
                        cell.text = re.sub(r'_{3,}', value, cell.text, count=1)
                    else:
                        cell.text = cell.text + ' ' + value
                    
                    return True
            
            return False
            
        except Exception as e:
            print(f"Error filling field {field['label']}: {str(e)}")
            return False


# Example usage
if __name__ == "__main__":
    filler = FormFiller()
    
    # Example extracted data
    extracted_data = {
        'dates': ['01/15/2024', 'March 20, 2024'],
        'emails': ['john.doe@email.com'],
        'phone_numbers': ['555-123-4567'],
        'numbers': ['123456789', '1234']
    }
    
    transcription = "My name is John Doe. I work in the Engineering department..."
    
    # Example usage with detected fields
    # fields = [...]  # From field detector
    # result = filler.fill_form(
    #     'form.docx',
    #     fields,
    #     extracted_data,
    #     transcription,
    #     'filled_form.docx'
    # )