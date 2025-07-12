# src/core/docx_creator.py
"""
Lesson 5: Creating Positioned DOCX Documents
Goal: Place text boxes EXACTLY where we found fields in the image

Key Concepts:
1. DOCX uses "points" and "inches" for positioning
2. We need to convert pixel coordinates to document coordinates
3. Text boxes can be positioned absolutely in DOCX
"""

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
import io

class DocxCreator:
    """Creates DOCX files with positioned text boxes matching the original form"""
    
    def __init__(self):
        # Standard US Letter size in inches
        self.page_width_inches = 8.5
        self.page_height_inches = 11.0
        
        # We'll calculate these based on the image
        self.pixels_per_inch_x = None
        self.pixels_per_inch_y = None
        
    def calculate_scaling(self, image_width, image_height):
        """
        Calculate how to convert pixels to inches
        Teaching: Images have pixels, documents have inches
        """
        print("\n📐 Calculating scaling factors...")
        
        # Assume the image represents a full page
        self.pixels_per_inch_x = image_width / self.page_width_inches
        self.pixels_per_inch_y = image_height / self.page_height_inches
        
        print(f"  Image size: {image_width} x {image_height} pixels")
        print(f"  Page size: {self.page_width_inches} x {self.page_height_inches} inches")
        print(f"  Scale: {self.pixels_per_inch_x:.1f} x {self.pixels_per_inch_y:.1f} pixels/inch")
        
    def pixels_to_inches(self, x_pixels, y_pixels):
        """Convert pixel coordinates to inches"""
        x_inches = x_pixels / self.pixels_per_inch_x
        y_inches = y_pixels / self.pixels_per_inch_y
        return x_inches, y_inches
    
    def add_positioned_textbox(self, doc, x_pixels, y_pixels, width_pixels, height_pixels, text=""):
        """
        Add a text box at a specific position
        Teaching: DOCX allows absolute positioning using frames
        """
        # Convert pixels to inches
        x_inches, y_inches = self.pixels_to_inches(x_pixels, y_pixels)
        width_inches = width_pixels / self.pixels_per_inch_x
        height_inches = height_pixels / self.pixels_per_inch_y
        
        # Create a paragraph
        paragraph = doc.add_paragraph()
        
        # Create a run (text within the paragraph)
        run = paragraph.add_run(text)
        
        # Style the text box
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        # Create frame for positioning (this is the tricky part!)
        # Frames allow absolute positioning in DOCX
        frame_props = parse_xml(f'''
            <w:framePr {nsdecls('w')}
                w:w="{int(width_inches * 1440)}"
                w:h="{int(height_inches * 1440)}"
                w:x="{int(x_inches * 1440)}"
                w:y="{int(y_inches * 1440)}"
                w:hAnchor="page"
                w:vAnchor="page"
                w:xAlign="left"
                w:yAlign="top"
                w:wrap="around"/>
        ''')
        # Note: 1440 = twips per inch (DOCX unit)
        
        paragraph._element.get_or_add_pPr().append(frame_props)
        
        # Add border to text box (so we can see it)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(0.1)
        
        return paragraph
    
    def add_checkbox(self, doc, x_pixels, y_pixels, size_pixels, checked=False):
        """
        Add a checkbox at specific position
        Teaching: We'll use special characters for checkboxes
        """
        # Checkbox characters: ☐ (empty) ☑ (checked)
        checkbox_char = "☑" if checked else "☐"
        
        # Add as positioned text
        para = self.add_positioned_textbox(
            doc, x_pixels, y_pixels, 
            size_pixels, size_pixels, 
            checkbox_char
        )
        
        # Make it look like a checkbox
        run = para.runs[0]
        run.font.size = Pt(12)
        
    def add_background_image(self, doc, image_path):
        """
        Add the original form as a background
        Teaching: This helps with alignment verification
        """
        # This is complex - we'll add it as a watermark
        section = doc.sections[0]
        
        # We'll implement this in the advanced version
        # For now, we'll work without background
        pass
    
    def create_form_replica(self, original_image, form_structure, output_path):
        """
        Main method: Create DOCX matching the original form
        """
        print("\n" + "="*60)
        print("📝 CREATING DOCX REPLICA")
        print("="*60)
        
        # Get image dimensions
        height, width = original_image.shape[:2]
        self.calculate_scaling(width, height)
        
        # Create new document
        doc = Document()
        
        # Set page margins to zero for full-page layout
        section = doc.sections[0]
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        
        # Add title (this will be removed in final version)
        title = doc.add_heading('Form Recreation Test', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Spacing
        
        # Process each detected element
        print("\n📦 Adding detected elements:")
        
        # Add text fields
        text_fields = [e for e in form_structure['all_elements'] 
                      if e.element_type in ['field', 'text_field']]
        
        print(f"\n  Adding {len(text_fields)} text fields...")
        for i, field in enumerate(text_fields):
            # Add a text box with sample text
            sample_text = f"Field {i+1}"
            self.add_positioned_textbox(
                doc, 
                field.x, field.y,
                field.width, field.height,
                sample_text
            )
        
        # Add checkboxes
        checkboxes = [e for e in form_structure['all_elements'] 
                     if e.element_type == 'checkbox']
        
        print(f"  Adding {len(checkboxes)} checkboxes...")
        for checkbox in checkboxes:
            self.add_checkbox(
                doc,
                checkbox.x, checkbox.y,
                checkbox.width,
                checked=False
            )
        
        # Save document
        doc.save(output_path)
        print(f"\n✅ Document saved to: {output_path}")
        
        return doc

# Let's also create a simpler version for learning
class SimpleDocxCreator:
    """
    Simplified version for understanding the basics
    This creates a form with relative positioning (easier to understand)
    """
    
    def create_simple_form(self, form_structure, output_path):
        """Create a simple version of the form"""
        print("\n📄 Creating simple DOCX form...")
        
        doc = Document()
        doc.add_heading('Recreated Form - Simple Version', 0)
        
        # Group elements by approximate Y position (rows)
        elements_by_row = {}
        
        for element in form_structure['all_elements']:
            # Round Y to nearest 30 pixels (group into rows)
            row_y = (element.y // 30) * 30
            
            if row_y not in elements_by_row:
                elements_by_row[row_y] = []
            
            elements_by_row[row_y].append(element)
        
        # Sort rows by Y position
        sorted_rows = sorted(elements_by_row.keys())
        
        # Create document content row by row
        for row_y in sorted_rows:
            elements = elements_by_row[row_y]
            
            # Sort elements in row by X position
            elements.sort(key=lambda e: e.x)
            
            # Create paragraph for this row
            para = doc.add_paragraph()
            
            for element in elements:
                if element.element_type == 'checkbox':
                    para.add_run('☐ ')
                elif element.element_type in ['field', 'text_field']:
                    # Add underlined space for field
                    run = para.add_run('_' * 20 + ' ')
                    run.font.underline = True
                else:
                    # Skip lines for now
                    continue
            
            # Add line break between rows
            if para.text.strip():  # Only if row has content
                para.add_run('\n')
        
        doc.save(output_path)
        print(f"✅ Simple form saved to: {output_path}")
        
        return doc

# Test script
if __name__ == "__main__":
    from pathlib import Path
    from image_preprocessor import ImagePreprocessor
    from form_structure_detector import FormStructureDetector
    
    print("🚀 Testing DOCX Creation...")
    
    # Load image and detect structure
    ticket_path = Path("../../data/sample_images/ticket.jpg")
    
    if ticket_path.exists():
        # Process image
        preprocessor = ImagePreprocessor()
        preprocessor.debug = False  # Turn off debug windows
        results = preprocessor.process_ticket_image(ticket_path)
        
        # Detect structure  
        detector = FormStructureDetector(debug=False)
        structure = detector.analyze_form_structure(results['binary'])
        
        # Create DOCX - Simple version first
        simple_creator = SimpleDocxCreator()
        simple_output = Path("../../data/output/form_simple.docx")
        simple_creator.create_simple_form(structure, simple_output)
        
        # Create DOCX - Advanced positioned version
        creator = DocxCreator()
        advanced_output = Path("../../data/output/form_advanced.docx")
        creator.create_form_replica(
            results['original'],
            structure,
            advanced_output
        )
        
        print("\n🎉 Success! Check the output folder for:")
        print(f"  1. {simple_output.name} - Simple version")
        print(f"  2. {advanced_output.name} - Advanced positioned version")
        
    else:
        print(f"⚠️ Image not found at: {ticket_path}")