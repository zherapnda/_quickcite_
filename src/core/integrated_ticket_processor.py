# src/core/integrated_ticket_processor.py
"""
Lesson 7: Complete Integration - Image + Audio = Filled Form
Goal: Extract info from audio and intelligently fill the detected fields

This combines:
1. Your existing audio transcription (Whisper)
2. OCR text extraction
3. Smart field matching
4. Final DOCX generation
"""

import cv2
import numpy as np
from pathlib import Path
import json
import re
from datetime import datetime
from typing import Dict, List, Tuple

# We'll import from your original project
import sys
sys.path.append('../../form_filler_ai/src/core')  # Adjust path as needed

try:
    from audio_transcriber import AudioTranscriber
    from form_filler import FormFiller
    AUDIO_AVAILABLE = True
except ImportError:
    print("⚠️ Audio modules not found. We'll create simplified versions.")
    AUDIO_AVAILABLE = False

class IntegratedTicketFiller:
    """
    The main class that brings everything together!
    Image → Structure → OCR → Audio → Filled Form
    """
    
    def __init__(self, whisper_model="base"):
        # Initialize all components
        from ocr_extractor import OCRExtractor, TicketProcessor
        from form_structure_detector import FormStructureDetector
        from image_preprocessor import ImagePreprocessor
        from docx_creator import DocxCreator
        
        self.preprocessor = ImagePreprocessor()
        self.structure_detector = FormStructureDetector(debug=False)
        self.ocr_extractor = OCRExtractor()
        self.docx_creator = DocxCreator()
        
        # Audio components
        if AUDIO_AVAILABLE:
            self.audio_transcriber = AudioTranscriber(model_size=whisper_model)
            self.form_filler = FormFiller()
        else:
            self.audio_transcriber = SimpleAudioTranscriber()
            
        self.debug = True
        
    def extract_info_from_conversation(self, audio_path):
        """
        Extract structured information from police-driver conversation
        Teaching: Real conversations are messy - we need smart extraction
        """
        print("\n🎙️ Processing audio recording...")
        
        # Transcribe audio
        if AUDIO_AVAILABLE:
            result = self.audio_transcriber.process_media_file(str(audio_path))
            transcription = result['text']
            
            # Extract entities
            entities = self.audio_transcriber.extract_entities(transcription)
        else:
            # Simplified version for testing
            transcription = self.simple_transcribe(audio_path)
            entities = self.simple_extract_entities(transcription)
        
        print(f"\n📝 Transcription preview:")
        print(f"{transcription[:200]}...")
        
        # Parse conversation for ticket-specific information
        ticket_data = self.parse_police_conversation(transcription)
        
        # Combine with extracted entities
        ticket_data['entities'] = entities
        ticket_data['full_transcription'] = transcription
        
        return ticket_data
    
    def parse_police_conversation(self, transcription):
        """
        Extract ticket information from natural conversation
        Teaching: Domain-specific parsing for police interactions
        """
        print("\n👮 Parsing police conversation...")
        
        data = {
            'officer_name': None,
            'badge_number': None,
            'violation': None,
            'speed': None,
            'speed_limit': None,
            'location': None,
            'driver_name': None,
            'license_number': None,
            'vehicle_info': None,
            'additional_notes': []
        }
        
        # Convert to lowercase for matching
        text_lower = transcription.lower()
        
        # Officer identification patterns
        officer_patterns = [
            r"(?:i'm|i am|this is)\s+officer\s+(\w+)",
            r"officer\s+(\w+)\s+(?:here|speaking)",
            r"badge\s+(?:number\s+)?(\d+)",
        ]
        
        for pattern in officer_patterns:
            match = re.search(pattern, text_lower)
            if match:
                if 'officer' in pattern:
                    data['officer_name'] = match.group(1).title()
                else:
                    data['badge_number'] = match.group(1)
        
        # Speed violation patterns
        speed_patterns = [
            r"(?:going|doing|clocked at)\s+(\d+)\s+(?:mph|miles)",
            r"(\d+)\s+in\s+a\s+(\d+)(?:\s+zone)?",
            r"speed\s+limit\s+(?:is\s+)?(\d+)",
        ]
        
        for pattern in speed_patterns:
            match = re.search(pattern, text_lower)
            if match:
                if len(match.groups()) == 2:
                    data['speed'] = match.group(1)
                    data['speed_limit'] = match.group(2)
                elif 'limit' in pattern:
                    data['speed_limit'] = match.group(1)
                else:
                    data['speed'] = match.group(1)
        
        # Violation type patterns
        violation_patterns = [
            r"(?:for|citing you for|violation is)\s+([^.]+?)(?:\.|,|$)",
            r"(speeding|running a (?:red light|stop sign)|illegal (?:turn|parking))",
            r"(failure to (?:stop|yield|signal))",
        ]
        
        for pattern in violation_patterns:
            match = re.search(pattern, text_lower)
            if match:
                data['violation'] = match.group(1).strip()
                break
        
        # Location patterns
        location_patterns = [
            r"(?:at|on|near)\s+(\w+\s+(?:street|avenue|road|boulevard|highway))",
            r"intersection of\s+(\w+\s+and\s+\w+)",
            r"mile marker\s+(\d+)",
        ]
        
        for pattern in location_patterns:
            match = re.search(pattern, text_lower)
            if match:
                data['location'] = match.group(1).title()
                break
        
        # Driver information
        if "license" in text_lower:
            # Look for license number pattern
            license_match = re.search(r"license\s+(?:number\s+)?([a-z0-9]+)", text_lower)
            if license_match:
                data['license_number'] = license_match.group(1).upper()
        
        # Print what we found
        print("\n📋 Extracted from conversation:")
        for key, value in data.items():
            if value and key != 'full_transcription' and key != 'entities':
                print(f"  - {key}: {value}")
        
        return data
    
    def match_audio_to_form_fields(self, audio_data, ocr_results, form_structure):
        """
        Intelligently match audio information to detected form fields
        Teaching: This is where AI makes smart decisions
        """
        print("\n🧩 Matching audio data to form fields...")
        
        # Get all text fields from structure
        text_fields = [e for e in form_structure['all_elements'] 
                      if e.element_type in ['field', 'text_field']]
        
        # Get labels from OCR
        labels = ocr_results.get('labels', [])
        
        # Create field mappings
        field_mappings = []
        
        for field in text_fields:
            # Find the closest label to this field
            closest_label = None
            min_distance = float('inf')
            
            for label in labels:
                # Calculate distance
                distance = np.sqrt((field.x - label.x)**2 + (field.y - label.y)**2)
                
                # Label should be to the left or above the field
                if (label.x <= field.x + field.width and 
                    label.y <= field.y + 20 and
                    distance < min_distance):
                    closest_label = label
                    min_distance = distance
            
            if closest_label:
                # Match label to audio data
                label_text = closest_label.text.lower()
                value = self.find_matching_value(label_text, audio_data)
                
                if value:
                    field_mappings.append({
                        'field': field,
                        'label': closest_label.text,
                        'value': value,
                        'confidence': 0.8
                    })
                    print(f"  Matched: {closest_label.text} → {value}")
        
        return field_mappings
    
    def find_matching_value(self, label_text, audio_data):
        """
        Find the best matching value for a label
        Teaching: Fuzzy matching and keyword association
        """
        label_lower = label_text.lower().strip(':')
        
        # Direct mappings
        mappings = {
            'officer': audio_data.get('officer_name'),
            'badge': audio_data.get('badge_number'),
            'speed': audio_data.get('speed'),
            'violation': audio_data.get('violation'),
            'location': audio_data.get('location'),
            'driver license': audio_data.get('license_number'),
            'name': audio_data.get('driver_name'),
        }
        
        # Check direct mappings
        for key, value in mappings.items():
            if key in label_lower and value:
                return value
        
        # Check entities
        entities = audio_data.get('entities', {})
        
        # Date fields
        if any(word in label_lower for word in ['date', 'when', 'time']):
            dates = entities.get('dates', [])
            if dates:
                return dates[0]
        
        # Phone fields
        if any(word in label_lower for word in ['phone', 'contact', 'tel']):
            phones = entities.get('phone_numbers', [])
            if phones:
                return phones[0]
        
        return None
    
    def create_final_document(self, image_path, form_structure, ocr_results, 
                            field_mappings, output_path):
        """
        Create the final DOCX with all information filled in
        """
        print("\n📄 Creating final document...")
        
        # Load original image for reference
        original_image = cv2.imread(str(image_path))
        
        # Create base document
        from docx import Document
        from docx.shared import Inches, Pt
        
        doc = Document()
        
        # Set up page
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Add title
        title = doc.add_heading('Traffic Citation - Processed', 0)
        doc.add_paragraph()
        
        # Calculate scaling
        height, width = original_image.shape[:2]
        self.docx_creator.calculate_scaling(width, height)
        
        # Add all OCR text first (as background text)
        for text_block in ocr_results['text_blocks']:
            # Skip if this is a label we're replacing
            is_replaced = any(
                mapping['label'].lower() in text_block.text.lower() 
                for mapping in field_mappings
            )
            
            if not is_replaced:
                # Add as positioned text
                self.docx_creator.add_positioned_textbox(
                    doc,
                    text_block.x, text_block.y,
                    text_block.width, text_block.height,
                    text_block.text
                )
        
        # Add filled fields
        for mapping in field_mappings:
            field = mapping['field']
            value = mapping['value']
            
            # Add filled text box
            para = self.docx_creator.add_positioned_textbox(
                doc,
                field.x, field.y,
                field.width, field.height,
                str(value)
            )
            
            # Highlight filled fields
            if para.runs:
                para.runs[0].font.bold = True
                para.runs[0].font.color.rgb = RGBColor(0, 0, 255)  # Blue
        
        # Add metadata
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        metadata = doc.add_paragraph()
        metadata.add_run("Processing Information:\n").bold = True
        metadata.add_run(f"Processed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        metadata.add_run(f"Fields filled: {len(field_mappings)}\n")
        metadata.add_run(f"Confidence: High\n")
        
        # Save document
        doc.save(output_path)
        print(f"✅ Document saved to: {output_path}")
        
        return doc
    
    def process_complete_ticket(self, image_path, audio_path, output_path):
        """
        The main method - processes everything!
        """
        print("\n" + "="*70)
        print("🚀 COMPLETE TICKET PROCESSING WITH AUDIO")
        print("="*70)
        
        # Step 1: Process image
        print("\n[Step 1/5] Processing image...")
        image_results = self.preprocessor.process_ticket_image(image_path)
        
        # Step 2: Detect structure
        print("\n[Step 2/5] Detecting form structure...")
        form_structure = self.structure_detector.analyze_form_structure(
            image_results['binary']
        )
        
        # Step 3: Extract text with OCR
        print("\n[Step 3/5] Extracting text with OCR...")
        ocr_results = self.ocr_extractor.extract_ticket_information(
            image_results['enhanced'],
            form_structure
        )
        
        # Step 4: Process audio
        print("\n[Step 4/5] Processing audio recording...")
        audio_data = self.extract_info_from_conversation(audio_path)
        
        # Step 5: Match and fill
        print("\n[Step 5/5] Matching audio to fields and creating document...")
        field_mappings = self.match_audio_to_form_fields(
            audio_data, ocr_results, form_structure
        )
        
        # Create final document
        final_doc = self.create_final_document(
            image_path, form_structure, ocr_results,
            field_mappings, output_path
        )
        
        # Summary
        print("\n" + "="*70)
        print("✅ PROCESSING COMPLETE!")
        print("="*70)
        print(f"\n📊 Summary:")
        print(f"  - Form fields detected: {len([e for e in form_structure['all_elements'] if 'field' in e.element_type])}")
        print(f"  - Text extracted: {len(ocr_results['text_blocks'])} blocks")
        print(f"  - Audio duration: {len(audio_data.get('full_transcription', ''))} characters")
        print(f"  - Fields filled: {len(field_mappings)}")
        print(f"  - Output: {output_path}")
        
        return {
            'success': True,
            'field_mappings': field_mappings,
            'audio_data': audio_data,
            'output_path': output_path
        }

# Simplified audio transcriber for testing
class SimpleAudioTranscriber:
    """Fallback audio transcriber if original modules not available"""
    
    def process_media_file(self, audio_path):
        """Simulate audio transcription"""
        # In real implementation, this would use Whisper
        mock_transcription = """
        This is Officer Johnson, badge number 5847. 
        I pulled you over for speeding. You were doing 45 in a 25 zone.
        This occurred at Main Street and 5th Avenue.
        Can I see your license please? 
        Alright Mr. Smith, I'm going to have to cite you for speeding today.
        The violation code is 4511.21. You have the option to appear in court
        or pay the fine online. The court date is set for next month.
        """
        
        return {
            'text': mock_transcription,
            'segments': [],
            'language': 'en'
        }
    
    def extract_entities(self, text):
        """Simple entity extraction"""
        import re
        
        entities = {
            'dates': re.findall(r'\b(?:next month|tomorrow|today)\b', text),
            'numbers': re.findall(r'\b\d+\b', text),
            'emails': [],
            'phone_numbers': []
        }
        
        return entities

# Test the complete system
if __name__ == "__main__":
    print("🎫 Integrated Ticket Processing System")
    print("="*50)
    
    # Set up paths
    ticket_image = Path("../../data/sample_images/ticket.jpg")
    audio_file = Path("../../data/sample_audio/officer_recording.mp3")
    output_file = Path("../../data/output/filled_ticket_complete.docx")
    
    # Check if files exist
    if not ticket_image.exists():
        print(f"⚠️ Please add ticket image to: {ticket_image}")
        print("\nFor testing, you can use your existing ticket image.")
    
    if not audio_file.exists():
        print(f"\n⚠️ Audio file not found at: {audio_file}")
        print("For testing, create a simple recording saying:")
        print('  "This is Officer Johnson, badge 5847."')
        print('  "You were speeding, doing 45 in a 25 zone."')
        print('  "This happened at Main Street."')
        print("\nOr we'll use mock data for testing.")
        
        # For testing without audio
        audio_file = None
    
    if ticket_image.exists():
        # Create processor
        processor = IntegratedTicketFiller(whisper_model="tiny")  # Use tiny for speed
        
        # Process ticket
        if audio_file and audio_file.exists():
            results = processor.process_complete_ticket(
                ticket_image, audio_file, output_file
            )
        else:
            print("\n📝 Using mock audio data for testing...")
            # Create mock audio file path
            mock_audio = Path("mock_audio.mp3")
            results = processor.process_complete_ticket(
                ticket_image, mock_audio, output_file
            )
        
        print("\n🎉 Success! Check your output folder!")
        print("\n💡 Next steps:")
        print("  1. Open the filled DOCX file")
        print("  2. Verify the audio information was filled correctly")
        print("  3. Check positioning and formatting")
        
        cv2.waitKey(0)
        cv2.destroyAllWindows()